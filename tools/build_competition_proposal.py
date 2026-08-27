from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_DOCX = DOCS / "資訊組_NutriGuard_構想提案書.docx"
DIAGRAM = ROOT / "tmp" / "proposal_review" / "nutriguard_flow.png"

FONT_NAME = "Microsoft JhengHei"
FONT_PATH = Path(r"C:\Windows\Fonts\msjh.ttc")
BOLD_FONT_PATH = Path(r"C:\Windows\Fonts\msjhbd.ttc")

BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "EAF2F8"
LIGHT_GREEN = "EAF6EF"
LIGHT_YELLOW = "FFF7E0"
LIGHT_RED = "FDECEC"
GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "B9C3CF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def style_cell_text(cell, bold=False, size=9.5, color=RGBColor(0, 0, 0), align=None) -> None:
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color


def set_run_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles["Heading 1"]
    elif level == 2:
        p.style = doc.styles["Heading 2"]
    else:
        p.style = doc.styles["Heading 3"]
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_para(doc: Document, text: str = "", bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_note_table(doc: Document, text: str, fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table, color="C8D5E3", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = text
    style_cell_text(cell, size=10.5)
    doc.add_paragraph()
    return table


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_table_width(table)
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, GRAY)
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style_cell_text(cell, bold=True, size=9.2, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row = table.add_row()
        set_cant_split(row)
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            style_cell_text(cell, size=8.6)
    doc.add_paragraph()
    return table


def add_label_table(doc: Document, rows: list[tuple[str, str]], widths=(2100, 7260)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_table_width(table)
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        row = table.rows[idx]
        set_cant_split(row)
        c0, c1 = row.cells
        c0.text = label
        c1.text = value
        set_cell_width(c0, widths[0])
        set_cell_width(c1, widths[1])
        set_cell_margins(c0)
        set_cell_margins(c1)
        set_cell_shading(c0, GRAY)
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style_cell_text(c0, bold=True, size=9.5, color=BLUE)
        style_cell_text(c1, size=9.5)
    doc.add_paragraph()
    return table


def add_tool_table(doc: Document):
    headers = ["工具類別", "工具名稱", "用途說明"]
    rows = [
        ["文字生成", "ChatGPT / Codex", "協助提案文字潤飾、章節重組、風險聲明與送件格式檢查；最終內容由團隊確認。"],
        ["圖像/影音生成", "未使用", "本提案書未使用生成式 AI 製作圖像或影音素材。"],
        ["程式碼輔助", "Codex", "協助 Expo 前端、Flask API、推薦邏輯、測試與文件整理等程式開發輔助。"],
        ["其他", "Gemini Vision", "作為作品功能的一部分，用於食物影像初判、候選食物名稱與份量描述；營養數值仍以資料庫查詢為主。"],
    ]
    return add_matrix(doc, headers, rows, [1700, 2100, 5560])


def draw_flow_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 860
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)

    font = ImageFont.truetype(str(FONT_PATH), 34)
    font_small = ImageFont.truetype(str(FONT_PATH), 25)
    font_bold = ImageFont.truetype(str(BOLD_FONT_PATH), 34)

    boxes = [
        (70, 110, 360, 250, "使用者輸入", "拍照、OCR、手動搜尋\n健康條件與過敏原"),
        (455, 110, 745, 250, "AI 初判", "Gemini Vision 產生\n食物候選與份量描述"),
        (840, 110, 1130, 250, "資料庫校準", "TFDA / 自訂食品庫\n換算營養素"),
        (1225, 110, 1515, 250, "安全過濾", "疾病規則與過敏原\n高鈉/高糖/風險提示"),
        (650, 470, 940, 610, "推薦排序", "剩餘熱量、營養匹配\n偏好回饋與 match_score"),
        (1035, 470, 1325, 610, "附近探索", "Google Places 真實店家\n營養資訊需再確認"),
        (1420, 470, 1710, 610, "輸出結果", "紀錄、警示、候選餐點\n與可解釋推薦理由"),
    ]

    def rounded_box(x1, y1, x2, y2, title, body, fill):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline="#8FA9C4", width=4)
        draw.text((x1 + 22, y1 + 20), title, fill="#163B5C", font=font_bold)
        draw.text((x1 + 22, y1 + 74), body, fill="#1E293B", font=font_small, spacing=8)

    fills = ["#EAF2F8", "#EAF6EF", "#FFF7E0", "#FDECEC", "#F4F6F9", "#EAF2F8", "#EAF6EF"]
    for box, fill in zip(boxes, fills):
        rounded_box(*box, fill=fill)

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="#4A657F", width=6)
        dx, dy = x2 - x1, y2 - y1
        if abs(dx) >= abs(dy):
            direction = 1 if dx > 0 else -1
            pts = [(x2, y2), (x2 - 24 * direction, y2 - 14), (x2 - 24 * direction, y2 + 14)]
        else:
            direction = 1 if dy > 0 else -1
            pts = [(x2, y2), (x2 - 14, y2 - 24 * direction), (x2 + 14, y2 - 24 * direction)]
        draw.polygon(pts, fill="#4A657F")

    arrow(360, 180, 455, 180)
    arrow(745, 180, 840, 180)
    arrow(1130, 180, 1225, 180)
    arrow(1370, 250, 840, 470)
    arrow(940, 540, 1035, 540)
    arrow(1325, 540, 1420, 540)
    arrow(795, 250, 795, 470)

    draw.text((70, 710), "可信機制：AI 僅做初判；營養數值優先由資料庫換算；低信心或缺資料時要求使用者確認。", fill="#334155", font=font)
    draw.text((70, 760), "責任邊界：本系統為健康管理輔助，不提供醫療診斷、治療建議或緊急過敏防護。", fill="#9B1C1C", font=font)
    img.save(path)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, before, after, color in [
        ("Heading 1", 15, 12, 6, BLUE),
        ("Heading 2", 12.5, 8, 4, BLUE),
        ("Heading 3", 11.5, 6, 3, RGBColor(31, 58, 95)),
    ]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NutriGuard 構想提案書")
    set_run_font(r, size=8, color=RGBColor(100, 116, 139))


def scrub_metadata(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "docProps/custom.xml":
                continue
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                data = (
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
                    'xmlns:dc="http://purl.org/dc/elements/1.1/" '
                    'xmlns:dcterms="http://purl.org/dc/terms/" '
                    'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
                    'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
                    "<dc:title>資訊組_NutriGuard_構想提案書</dc:title>"
                    "<dc:creator></dc:creator>"
                    "<cp:lastModifiedBy></cp:lastModifiedBy>"
                    "</cp:coreProperties>"
                ).encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def build() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    draw_flow_diagram(DIAGRAM)

    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("2026 銘傳大學 AI 應用創意大賽：構想提案書")
    set_run_font(r, size=18, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("NutriGuard：個人化健康飲食 AI 推薦與即時辨識系統")
    set_run_font(r, size=13, bold=True)

    add_label_table(
        doc,
        [
            ("作品名稱", "NutriGuard：個人化健康飲食 AI 推薦與即時辨識系統"),
            ("參賽組別", "資訊組"),
            ("競賽主軸", "AI 生活創新"),
            ("提案定位", "以食物辨識、營養資料查詢、疾病/過敏原風險提醒與附近店家探索，整合成可展示的日常健康飲食決策流程。"),
        ],
    )

    add_note_table(
        doc,
        "匿名原則：本文件不揭露參賽者姓名或任何足以辨識個人身分之資訊。",
        fill=LIGHT_YELLOW,
    )

    add_heading(doc, "2、應用介紹（創意核心與痛點解決）", 1)
    add_heading(doc, "(1) 問題背景與痛點分析", 2)
    add_para(
        doc,
        "外食、便利餐點與包裝食品已成為許多人的日常選擇，但熱量、鈉含量、糖分、過敏原與慢性病飲食禁忌往往分散在不同標示或資料來源中。對外食族、體重管理者、三高與慢性病風險族群而言，手動查詢食品、估算份量與判斷風險成本過高，導致飲食紀錄難以持續，健康建議也常與真實用餐場景脫節。",
    )
    add_bullet(doc, "目標對象：校園外食族、超商/便當/餐飲外帶使用者，以及需留意高鈉、高糖、高脂、蛋白質或過敏原風險的使用者。")
    add_bullet(doc, "主要痛點：手動紀錄繁瑣、外食營養資訊不足、疾病/過敏原提醒不即時、附近店家推薦與個人健康限制缺乏連結。")

    add_heading(doc, "(2) 作品核心創意", 2)
    add_para(
        doc,
        "NutriGuard 的核心創意是將「AI 食物初判」、「可信營養資料換算」、「疾病/過敏原安全過濾」與「附近店家探索」整合成同一個日常決策流程。使用者可透過拍照、營養標示 OCR 或手動搜尋建立飲食紀錄，系統再依健康條件、今日剩餘熱量、營養匹配與近期偏好提供可解釋推薦。",
    )
    add_matrix(
        doc,
        ["創新維度", "常見作法", "NutriGuard 作法", "差異與效益"],
        [
            [
                "辨識與記錄",
                "使用者手動搜尋食品名稱，並自行估算份量與營養素。",
                "Gemini Vision 先做食物語意初判，再由 TFDA/自訂食品資料庫換算營養；低信心時導向 OCR 或手動確認。",
                "降低查表與重複輸入負擔，同時避免把 AI 猜測直接當成營養真值。",
            ],
            [
                "疾病與過敏提醒",
                "使用者自行查詢疾病禁忌，外食時容易忽略高風險成分。",
                "依 disease_rules.json 與過敏原設定提示高鈉、高糖、高脂、蛋白質或過敏風險。",
                "讓健康雷區在介面中可見，但明確定位為輔助資訊，不取代醫療或營養專業判斷。",
            ],
            [
                "推薦與地圖",
                "固定卡路里建議或一般店家搜尋，與使用者所在地與健康條件連結不足。",
                "以規則型雙軌推薦排序日常候選食品，並用 Google Places 搜尋附近真實店家。",
                "先找到可前往的場景，再由使用者依店家標示、掃描或手動搜尋確認實際餐點。",
            ],
        ],
        [1450, 2350, 3000, 2560],
    )

    add_heading(doc, "3、操作方式（運作邏輯與技術實作）", 1)
    add_heading(doc, "(1) 系統運作邏輯", 2)
    add_para(doc, "系統流程分為輸入、AI 初判、資料庫校準、安全過濾、推薦排序與輸出六個階段。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(DIAGRAM), width=Inches(6.25))
    add_para(doc, "圖 1. NutriGuard 系統運作流程與可信設計。")

    add_heading(doc, "(2) AI 技術實作流程", 2)
    add_label_table(
        doc,
        [
            ("前端與使用流程", "Expo React Native 實作 Web/iOS/Android 可延伸介面，支援拍照、OCR、手動搜尋、飲食紀錄、推薦與地圖探索。"),
            ("後端與資料服務", "Flask API 部署於 Render，資料以 Supabase Postgres/Auth 為主，保留 MongoDB / In-memory fallback 作為開發與備援路徑。"),
            ("食物辨識", "Gemini Vision 產生可能食物名稱、份量描述、同義候選與信心程度；低信心或無資料庫對應時，要求使用者手動確認。"),
            ("營養 grounding", "熱量、三大營養素、鈉、纖維等數值優先由 TFDA 台灣食品營養資料庫與自訂食品資料換算。"),
            ("安全與推薦", "Safety Guard 依疾病規則與過敏原設定篩除或提示高風險項目；推薦分數整合剩餘熱量、NutritionMatch、PreferenceScore 與 FeedbackAdjustment。"),
            ("地圖探索", "Google Places 提供真實店家位置、評分、營業狀態與距離；因其不提供可靠營養與過敏原資料，系統標示需由現場標示、掃描或手動搜尋確認。"),
        ],
        widths=(2000, 7360),
    )
    add_note_table(
        doc,
        "可信設計：AI 僅作初判與摘要；營養數值以資料庫為主；低信心、缺資料或高風險情境需由使用者依標示、店家資訊或專業建議確認。",
        fill=LIGHT_GREEN,
    )

    add_heading(doc, "4、預期效果（社會影響力、永續價值與成效）", 1)
    add_heading(doc, "(1) 社會影響力", 2)
    add_para(
        doc,
        "NutriGuard 以日常飲食決策為切入點，讓外食族與慢性病風險族群更容易看見食品營養、疾病禁忌與過敏原風險。作品的競賽展示重點不是取代營養師，而是把食物辨識、營養資料查詢、疾病風險提醒與附近店家探索整合成可操作的輔助流程，降低使用者在外食場景中的資訊落差。",
    )
    add_heading(doc, "(2) 永續價值（SDGs）", 2)
    add_bullet(doc, "SDG 3 良好健康與福祉：透過日常飲食風險提醒，協助慢性病風險族群進行預防性健康管理。")
    add_bullet(doc, "SDG 12 責任消費與生產：協助使用者理解包裝食品與外食營養資訊，做出更透明、可追蹤的消費選擇。")
    add_heading(doc, "(3) 實踐成效預測", 2)
    add_matrix(
        doc,
        ["評估項目", "衡量方式", "MVP 目標"],
        [
            ["紀錄便利性", "以 10-20 名校園外食使用者測試單筆飲食紀錄時間與操作步驟。", "相較純手動輸入，降低重複查表與輸入步驟。"],
            ["AI 初判可信度", "記錄 AI 初判後需手動修正比例、低信心轉手動確認比例。", "建立可追蹤的修正資料，作為後續模型與資料庫改善依據。"],
            ["風險提示理解度", "以問卷或訪談確認高鈉、高糖、過敏原與疾病禁忌提醒是否易懂。", "讓高風險資訊能在介面上被明確辨識，降低誤解。"],
            ["推薦可用性", "觀察使用者是否採納、略過或標記不喜歡推薦項目。", "累積 preference_score 與回饋資料，改善規則型排序。"],
        ],
        [1900, 4300, 3160],
    )
    add_heading(doc, "(4) 風險、倫理與責任邊界", 2)
    add_para(
        doc,
        "使用者健康條件、過敏原與飲食紀錄屬敏感健康資料，系統應採明確同意、最小化蒐集、刪除權限與去識別化分析。NutriGuard 不提供醫療診斷、治療建議、疾病控制保證或緊急過敏防護；所有 AI 辨識、店家摘要與推薦結果皆需由使用者依實際標示、店家資訊與專業建議確認。",
    )

    add_heading(doc, "5、生成式 AI 工具使用清單", 1)
    add_para(doc, "本團隊於製作過程與作品功能中使用之生成式 AI 工具如下：")
    add_tool_table(doc)
    add_note_table(
        doc,
        "規範聲明：本團隊保證清單中不包含 DeepSeek 及任何由中國大陸開發之 AI 工具（如：文心一言、通義千問等）。若經查證屬實，願接受取消參賽及得獎資格之處分。",
        fill=LIGHT_YELLOW,
    )

    add_heading(doc, "6、版權宣告", 1)
    add_label_table(
        doc,
        [
            ("(1) 原創保證", "本作品之系統設計、程式實作、提案文字與展示內容均由本團隊整理與製作，或已依規定取得合法授權。"),
            ("(2) 素材授權說明", "本作品若有使用具著作權之素材，相關合法授權證明文件、原始來源截圖或授權紀錄將另行整合於「其他證明文件」檔案中上傳，以利查核。"),
            ("(3) 責任承擔聲明", "本團隊保證無侵權行為；若有不實，願自負法律責任並接受取消參賽及得獎資格。"),
        ],
        widths=(2500, 6860),
    )

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "資訊組_NutriGuard_構想提案書"
    doc.core_properties.comments = ""
    doc.save(OUT_DOCX)
    scrub_metadata(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
