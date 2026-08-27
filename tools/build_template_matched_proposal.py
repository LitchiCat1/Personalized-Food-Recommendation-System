from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_DOCX = DOCS / "資訊組_NutriGuard_構想提案書.docx"

FONT = "Microsoft JhengHei"
FONT_SCALE = 1.12


def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size * FONT_SCALE)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para(p, left=0, first=0, before=0, after=2, line=1.12):
    p.paragraph_format.left_indent = Inches(left)
    p.paragraph_format.first_line_indent = Inches(first)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def add_text(doc, text="", left=0, first=0, size=12, bold=False, color=None, before=0, after=2, align=None):
    p = doc.add_paragraph()
    set_para(p, left=left, first=first, before=before, after=after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, left=0.48, size=12, color=None, before=0, after=1):
    p = doc.add_paragraph()
    set_para(p, left=left, first=-0.18, before=before, after=after)
    r1 = p.add_run("●")
    set_font(r1, size=size, color=color)
    r2 = p.add_run("  " + text)
    set_font(r2, size=size, color=color)
    return p


def set_cell_text(cell, text, size=12, bold=False):
    cell.text = text
    for p in cell.paragraphs:
        set_para(p, after=0, line=1.05)
        for run in p.runs:
            set_font(run, size=size, bold=bold)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, size="10", color="000000"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_width(table, width=9360, indent=720):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def add_tool_table(doc):
    rows = [
        ["工具類別", "工具名稱", "用途說明"],
        ["文字生成", "OpenCode / Codex", "協助提案文字整理、格式檢查與潤飾；最終內容由團隊確認。"],
        ["圖像/影音生成", "未使用", "本提案書未使用生成式 AI 製作圖像或影音素材。"],
        ["程式碼輔助", "Codex", "協助 Expo 前端、Flask API、推薦邏輯、測試與文件整理等程式開發輔助。"],
        ["其他", "Gemini Vision", "作為作品功能的一部分，用於食物影像初判、候選食物名稱與份量描述；營養數值仍以資料庫查詢為主。"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    set_table_width(table, width=9360, indent=720)
    set_table_borders(table, size="10")
    widths = [2300, 3300, 3760]
    for ri, row_data in enumerate(rows):
        for ci, text in enumerate(row_data):
            cell = table.cell(ri, ci)
            set_cell_text(cell, text, size=12, bold=(ri == 0))
            set_cell_width(cell, widths[ci])
            set_cell_margins(cell)
    return table


def scrub_metadata(path: Path):
    tmp = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "docProps/custom.xml":
                continue
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                data = (
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
                    'xmlns:dc="http://purl.org/dc/elements/1.1/" '
                    'xmlns:dcterms="http://purl.org/dc/terms/" '
                    'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
                    'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
                    "<dc:title>資訊組_NutriGuard_構想提案書</dc:title>"
                    "<dc:creator></dc:creator>"
                    "<cp:lastModifiedBy></cp:lastModifiedBy>"
                    "</cp:coreProperties>"
                ).encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1.12)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(1.05)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12 * FONT_SCALE)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.12

    p = add_text(doc, "2026 銘傳大學 AI 應用創意大賽：構想提案書", size=18, bold=True, before=18, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.line_spacing = 1

    add_text(doc, "（本文件為提案書格式範本，參賽團隊可自行下載並修改使用）", left=0.05, size=11.5, after=0)
    add_bullet(doc, "建議頁數: 5-10 頁，", left=0.25, size=11.5, after=0)
    p = doc.add_paragraph()
    set_para(p, left=0.25, first=-0.18, after=0)
    r = p.add_run("●  ")
    set_font(r, size=11.5)
    r = p.add_run("檔名: 組別_作品名稱_構想提案書.pdf (範例: 跨領域組_智慧綠校園_構想提案書.pdf)")
    set_font(r, size=11.5)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p = doc.add_paragraph()
    set_para(p, left=0.25, first=-0.18, after=14)
    r = p.add_run("●  ")
    set_font(r, size=11.5)
    r = p.add_run("重要提醒: 匿名原則本文件內嚴禁揭露參賽者姓名、系級、學號或足以辨識身份之資訊。違反\n   此規定者，主辦單位將予以扣分。")
    set_font(r, size=11.5, color=RGBColor(255, 0, 0))

    add_text(doc, "1、  作品基本資料作品名稱：NutriGuard：個人化健康飲食 AI 推薦與即時辨識系統", size=12.5, bold=True, after=0)
    add_text(doc, "（請輸入與報名表一致的名稱）", left=2.45, size=11.5, after=6)
    add_text(doc, "(1)  參賽組別： ■ 資訊組 / □ 非資訊組 / □ 跨領域組", left=0.28, size=12, after=2)
    add_text(doc, "(2)  競賽主軸： □ AI 永續 / ■ AI 生活創新", left=0.28, size=12, after=16)

    add_text(doc, "2、  應用介紹（創意核心與痛點解決）", size=12.5, bold=True, after=4)
    add_text(doc, "(1)  問題背景與痛點分析：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "外食、便利餐點與包裝食品已成為許多人的日常選擇，但熱量、鈉含量、糖分、過敏原與慢性病飲食禁忌往往分散在不同標示或資料來源中，導致飲食紀錄與風險判斷成本過高。", left=0.73, size=11.5)
    add_bullet(doc, "多數飲食紀錄 App 要求使用者手動搜尋食品名稱、估算份量與輸入營養數值，流程繁瑣且容易中斷，外食族在真實用餐場景中很難長期維持。", left=0.73, size=11.5)
    add_bullet(doc, "慢性病風險族群與嚴重過敏者在外食時常面臨資訊不足，例如食品成分標示不清、菜單缺少營養資料，或沒有即時提醒高鈉、高糖、高脂、蛋白質與過敏原風險。", left=0.73, size=11.5)
    add_bullet(doc, "目標對象為校園外食族、超商/便當/餐飲外帶使用者，以及需留意高鈉、高糖、高脂、蛋白質或過敏原風險的使用者。", left=0.73, size=11.5, after=5)
    add_text(doc, "(2)  作品核心創意：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "NutriGuard 將 AI 食物初判、可信營養資料換算、疾病/過敏原安全過濾與附近店家探索整合為同一個日常健康飲食決策流程。", left=0.73, size=11.5)
    add_bullet(doc, "與一般手動飲食紀錄 App 不同，本作品以 Gemini Vision 進行食物語意初判，再由 TFDA/自訂食品資料庫換算營養，低信心或缺資料時要求使用者以 OCR 或手動搜尋確認。", left=0.73, size=11.5)
    add_bullet(doc, "推薦邏輯採雙軌設計：第一軌先依疾病規則與過敏原設定處理安全風險，第二軌再依剩餘熱量、營養匹配、近期飲食紀錄與使用者回饋進行排序。", left=0.73, size=11.5)
    add_bullet(doc, "附近店家探索不是直接宣稱餐點健康，而是先協助使用者找到真實可前往店家，再提醒使用者依店家標示、掃描或手動搜尋確認實際餐點營養與過敏原資訊。", left=0.73, size=11.5, after=5)
    add_text(doc, "(3)  跨學院/跨領域合作成效說明（加分項目）：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "無。", left=0.73, size=11.5, after=14)

    add_text(doc, "3、  操作方式（運作邏輯與技術實作）", size=12.5, bold=True, after=4)
    add_text(doc, "(1)  系統運作邏輯：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "使用者拍照、營養標示 OCR 或手動搜尋建立輸入；系統以 Gemini Vision 初判食物名稱與份量，再由資料庫換算營養，接著依疾病規則、過敏原與今日剩餘熱量產生風險提醒與候選推薦。", left=0.73, size=11.5)
    add_bullet(doc, "日常紀錄流程為：影像/文字輸入 → 食物候選名稱與份量描述 → TFDA 或自訂食品資料庫比對 → 熱量、三大營養素、鈉、纖維等數值換算 → 疾病與過敏原規則檢查 → 使用者確認後寫入飲食紀錄。", left=0.73, size=11.5)
    add_bullet(doc, "推薦流程為：先排除或提示高風險項目，再依今日剩餘熱量、營養匹配、近期飲食偏好與推薦回饋計算候選項目的排序分數，並提供可解釋的推薦理由。", left=0.73, size=11.5)
    add_bullet(doc, "附近店家探索由 Google Places 提供真實位置、評分、營業狀態與距離；店家菜單、營養與過敏原資料仍需由使用者依現場標示、掃描或手動搜尋確認。", left=0.73, size=11.5, after=5)
    add_text(doc, "(2)  AI 技術實作流程：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "食物辨識採 Gemini Vision 產生候選食物名稱、份量描述、同義候選與信心程度；營養數值優先由 TFDA 台灣食品營養資料庫與自訂食品資料換算。", left=0.73, size=11.5)
    add_bullet(doc, "前端以 Expo React Native 建立 Web/iOS/Android 可延伸介面，後端以 Flask API 提供掃描、紀錄、推薦與地圖相關服務，資料以 Supabase Postgres/Auth 管理。", left=0.73, size=11.5)
    add_bullet(doc, "Safety Guard 會讀取疾病規則與過敏原設定，針對高血壓、糖尿病、高血脂、痛風、慢性腎臟病等情境提供高鈉、高糖、高脂、蛋白質或過敏風險提示。", left=0.73, size=11.5)
    add_bullet(doc, "推薦邏輯採可解釋的規則型雙軌設計：Safety Guard 先處理疾病與過敏原風險，再依剩餘熱量、營養匹配、近期飲食偏好與推薦回饋排序。此作法適合 MVP 展示，且能降低直接依賴生成式 AI 輸出精準營養數字的風險。", left=0.73, size=11.5, after=14)

    add_text(doc, "4、  預期效果（社會影響力、永續價值與成效）", size=12.5, bold=True, after=4)
    add_text(doc, "(1)  社會影響力：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "NutriGuard 讓外食族與慢性病風險族群更容易看見食品營養、疾病禁忌與過敏原風險，降低日常外食場景中的資訊落差。", left=0.73, size=11.5)
    add_bullet(doc, "本作品定位為健康管理輔助工具，不提供醫療診斷、治療建議、疾病控制保證或緊急過敏防護；所有 AI 辨識、店家摘要與推薦結果皆需由使用者依實際標示、店家資訊與專業建議確認。", left=0.73, size=11.5, after=4)
    add_text(doc, "(2)  永續價值（SDGs）：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "對應 SDG 3 良好健康與福祉，透過日常飲食風險提醒協助預防性健康管理，讓慢性病風險族群更容易在日常飲食中察覺高風險選擇。", left=0.73, size=11.5)
    add_bullet(doc, "對應 SDG 12 責任消費與生產，協助使用者理解包裝食品與外食營養資訊，讓消費選擇更透明、可追蹤。", left=0.73, size=11.5, after=4)
    add_text(doc, "(3)  實踐成效預測：", left=0.28, size=12, bold=True, after=1)
    add_bullet(doc, "MVP 階段預計以校園外食使用者測試單筆飲食紀錄時間、AI 初判後需手動修正比例、疾病/過敏提醒理解度與推薦採納回饋，作為後續資料庫與推薦邏輯改善依據。", left=0.73, size=11.5)
    add_bullet(doc, "預期觀察指標包含：相較純手動輸入是否減少查表與重複輸入步驟、風險提示是否能被使用者理解、以及推薦排序是否能依採納/略過/不喜歡回饋逐步改善。", left=0.73, size=11.5, after=14)

    add_text(doc, "5、  生成式 AI 工具使用清單", size=12.5, bold=True, after=1)
    add_text(doc, "請詳列製作過程中使用的所有生成式 AI 工具。", left=0.28, size=11.5, after=3)
    add_tool_table(doc)
    p = doc.add_paragraph()
    set_para(p, left=0.28, before=2, after=18)
    r = p.add_run("規範聲明: 本團隊保證清單中不包含 DeepSeek 及任何由中國大陸開發之\nAI 工具（如: 文心一言、通義千問等）。若經查證屬實，願接受取消參賽及得\n獎資格之處分。")
    set_font(r, size=12, bold=True, italic=True)

    add_text(doc, "6、  版權宣告", size=12.5, bold=True, after=2)
    add_text(doc, "(1)  原創保證: 本作品素材均為本團隊原創，或已依規定取得合法授權。", size=12, after=1)
    add_text(doc, "(2)  素材授權說明: 本作品若有使用具著作權之素材，相關合法授權證明文件\n     （版權授權書、原始來源截圖等）已另行整合於「其他證明文件」檔案中上傳，\n     以利查核。", size=12, after=1)
    add_text(doc, "(3)  責任承擔聲明: 本團隊保證無侵權行為，若有不實願自負法律責任並取消參\n     賽及得獎資格。", size=12, after=1)

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.core_properties.title = "資訊組_NutriGuard_構想提案書"
    doc.save(OUT_DOCX)
    scrub_metadata(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
